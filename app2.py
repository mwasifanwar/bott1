import streamlit as st
import openai
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
import os
from meeting_note_taker import *  # Import the new file
from financial_modeling import financial_modeling_page 

# Set your OpenAI API key
openai.api_key = os.getenv("OPENAI_API_KEY")


# Page configuration
st.set_page_config(
    page_title="Business Document Generator",
    page_icon=":chart_with_upwards_trend:",
    layout="wide",
)

# Sidebar navigation (existing code)
with st.sidebar:
    st.markdown("<h3 style='text-align: center;'>Wasif Muhammad</h3>", unsafe_allow_html=True)
    st.markdown("---")
    if st.button("Home", key="home_button"):
        st.session_state.page = "home"
    if st.button("Comprehensive Financial Modeling", key="financial_model_button"):
        st.session_state.page = "financial_model"
    if st.button("Meeting Note Taker", key="meeting_note_taker_button"):
        st.session_state.page = "meeting_note_taker"
    st.button("Others", key="others_button", help="Explore other options")


# Main content switcher
if "page" not in st.session_state:
    st.session_state.page = "home"

if st.session_state.page == "home":
    st.title("Welcome to the Business Tool")
    
    # Main content
    st.title("Use **ME** to create your **Business Documents**")
    st.markdown("### Enter Business Information")

    # Document type selection
    document_type = st.selectbox(
        "Select Document Type",
        ["Business Plan", "Feasibility Study", "Application Form", "Pitch Deck"]
    )

    # Create a 2-column layout for the input fields
    col1, col2, col3 = st.columns(3)

    with col1:
        language = st.selectbox("Language", ["UK English", "US English"])
        writing_person = st.selectbox("Writing Person", ["1st Person", "3rd Person"])
        writing_style = st.selectbox("Writing Style", ["Formal", "Informal"])

    with col2:
        document_length = st.selectbox("Length of Document", ["Short", "Long"])
        template = st.selectbox("Template", ["Standard", "IDC", "NEF", "Custom"])
        business_type = st.selectbox("Business Type", ["Start-up", "Expansion", "Acquisition"])

    with col3:
        bee_level = st.selectbox("BEE Level", ["Level 1", "Level 2", "Level 3", "Level 4"])
        directors = st.text_input("Directors/Shareholders", placeholder="Enter names")
        staffing = st.text_input("Staffing Compliment", placeholder="Enter staffing details")

    col4, col5, col6 = st.columns(3)

    with col4:
        funding_amount = st.text_input("Funding Amount", placeholder="Enter amount")

    with col5:
        grant_amount = st.text_input("Grant Amount", placeholder="Enter amount")

    with col6:
        finance_term = st.text_input("Finance Term", placeholder="Enter term")

    # Single column for the business overview and generate button
    business_overview = st.text_area("Your Business Overview...", placeholder="Enter a brief overview of your business")
    st.markdown("<br>", unsafe_allow_html=True)  # Add some space

    # Initialize session state to retain the generated plan
    if "generated_plan" not in st.session_state:
        st.session_state["generated_plan"] = None
    if "edited_plan" not in st.session_state:
        st.session_state["edited_plan"] = None

    # Function to replace placeholders in text with actual user inputs
    def replace_placeholders(text, context):
        for key, value in context.items():
            text = text.replace(f"[{key}]", value)
        return text

    # Function to save a chart as an image
    def save_chart_as_image(chart_type, filename):
        fig, ax = plt.subplots()

        if chart_type == "Market Share Over Time":
            ax.plot(["Year 1", "Year 2", "Year 3", "Year 4", "Year 5"], [10, 20, 30, 45, 60], marker='o')
            ax.set_xlabel('Years')
            ax.set_ylabel('Market Share (%)')
            
        elif chart_type == "Competitive Analysis":
            categories = ['Product Quality', 'Market Share', 'Innovation', 'Customer Service', 'Price']
            competitor1 = [8, 7, 9, 6, 5]
            competitor2 = [6, 8, 7, 9, 6]
            competitor3 = [7, 6, 8, 7, 8]
            N = len(categories)
            angles = [n / float(N) * 2 * 3.14159 for n in range(N)]
            angles += angles[:1]
            ax = plt.subplot(111, polar=True)
            plt.xticks(angles[:-1], categories)
            ax.set_rlabel_position(0)
            plt.yticks([2, 4, 6, 8, 10], ["2", "4", "6", "8", "10"], color="grey", size=7)
            plt.ylim(0, 10)
            ax.plot(angles, competitor1 + competitor1[:1], linewidth=1, linestyle='solid', label="Competitor 1")
            ax.plot(angles, competitor2 + competitor2[:1], linewidth=1, linestyle='solid', label="Competitor 2")
            ax.plot(angles, competitor3 + competitor3[:1], linewidth=1, linestyle='solid', label="Competitor 3")
            ax.fill(angles, competitor1 + competitor1[:1], 'b', alpha=0.1)
            ax.fill(angles, competitor2 + competitor2[:1], 'r', alpha=0.1)
            ax.fill(angles, competitor3 + competitor3[:1], 'g', alpha=0.1)
            
        elif chart_type == "Milestone Timeline":
            milestones = ['Setup', 'R&D', 'Launch', 'Market Expansion', 'Consolidation']
            timeline = [2, 6, 12, 24, 36]  # months
            ax.barh(milestones, timeline, color='orange')
            ax.set_xlabel('Months')
            
        elif chart_type == "Revenue vs Expenses":
            years = ["Year 1", "Year 2", "Year 3", "Year 4", "Year 5"]
            revenue = [200, 400, 600, 800, 1000]
            expenses = [150, 300, 500, 700, 850]
            ax.plot(years, revenue, label='Revenue', marker='o')
            ax.plot(years, expenses, label='Expenses', marker='o', linestyle='--')
            ax.set_xlabel('Years')
            ax.set_ylabel('Amount (£)')
            ax.legend()

        elif chart_type == "Cash Flow Forecast":
            months = [f'Month {i}' for i in range(1, 13)]
            cash_inflow = [1000, 1500, 1800, 2000, 2200, 2500, 2700, 2900, 3000, 3200, 3400, 3500]
            cash_outflow = [800, 900, 1000, 1200, 1300, 1400, 1500, 1600, 1700, 1800, 1900, 2000]
            cash_balance = [inflow - outflow for inflow, outflow in zip(cash_inflow, cash_outflow)]
            ax.plot(months, cash_inflow, label='Cash Inflow', marker='o')
            ax.plot(months, cash_outflow, label='Cash Outflow', marker='o')
            ax.plot(months, cash_balance, label='Cash Balance', marker='o', linestyle='--')
            ax.set_xlabel('Months')
            ax.set_ylabel('Amount (£)')
            ax.legend()

        elif chart_type == "Organizational Structure":
            positions = ["CEO", "CTO", "CFO", "COO", "CMO"]
            levels = [1, 2, 2, 3, 3]
            ax.barh(positions, levels, color='teal')
            ax.set_xlabel('Management Level')

        elif chart_type == "Product Development Roadmap":
            stages = ['Concept', 'Development', 'Testing', 'Launch', 'Post-Launch']
            months_to_complete = [2, 5, 3, 2, 6]
            ax.bar(stages, months_to_complete, color='purple')
            ax.set_xlabel('Development Stages')
            ax.set_ylabel('Time (Months)')

        plt.savefig(filename)
        plt.close()

    # Function to generate financial data tables
    def generate_financial_tables():
        revenue_data = ["£200,000", "£400,000", "£800,000", "£1,600,000", "£3,200,000"]
        expenses_data = ["£100,000", "£150,000", "£200,000", "£250,000", "£300,000"]
        net_profit_data = [str(int(rev.replace("£", "").replace(",", "")) - int(exp.replace("£", "").replace(",", ""))) for rev, exp in zip(revenue_data, expenses_data)]

        revenue_table = f"""
        | Year | Revenue (£) |
        |------|-------------|
        | 1    | {revenue_data[0]} |
        | 2    | {revenue_data[1]} |
        | 3    | {revenue_data[2]} |
        | 4    | {revenue_data[3]} |
        | 5    | {revenue_data[4]} |
        """

        operating_expenses_table = f"""
        | Year | Operating Expenses (£) |
        |------|------------------------|
        | 1    | {expenses_data[0]} |
        | 2    | {expenses_data[1]} |
        | 3    | {expenses_data[2]} |
        | 4    | {expenses_data[3]} |
        | 5    | {expenses_data[4]} |
        """

        net_profit_table = f"""
        | Year | Net Profit (£) |
        |------|----------------|
        | 1    | £{net_profit_data[0]} |
        | 2    | £{net_profit_data[1]} |
        | 3    | £{net_profit_data[2]} |
        | 4    | £{net_profit_data[3]} |
        | 5    | £{net_profit_data[4]} |
        """

        return revenue_table, operating_expenses_table, net_profit_table

    # Function to convert the generated business plan to DOCX format
    def convert_to_docx(business_plan, charts_to_generate):
        doc = Document()
        doc.add_heading('Business Document', 0)
        doc.add_paragraph(business_plan)

        for chart_type, title in charts_to_generate:
            doc.add_heading(title, level=1)
            # Save the chart as an image and insert it into the DOCX
            chart_filename = f"{chart_type.replace(' ', '_')}.png"
            save_chart_as_image(chart_type, chart_filename)
            doc.add_picture(chart_filename, width=Inches(5))  # Correct usage of Inches
            os.remove(chart_filename)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    # Custom PDF class to handle encoding
    class CustomPDF(FPDF):
        def __init__(self):
            super().__init__()
            self.set_font("Arial", size=12)
            self.add_page()

        def header(self):
            self.set_font("Arial", 'B', 12)
            self.cell(0, 10, 'Business Document', align='C', ln=True)

        def footer(self):
            self.set_y(-15)
            self.set_font("Arial", 'I', 8)
            self.cell(0, 10, f'Page {self.page_no()}', align='C')

    # Function to convert the generated business plan to PDF format
    def convert_to_pdf(business_plan, charts_to_generate):
        pdf = CustomPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, business_plan)

        for chart_type, title in charts_to_generate:
            pdf.add_page()
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(200, 10, txt=title, ln=True, align='C')

            # Save the chart as an image and insert it into the PDF
            chart_filename = f"{chart_type.replace(' ', '_')}.png"
            save_chart_as_image(chart_type, chart_filename)
            pdf.image(chart_filename, x=10, w=pdf.w - 20)
            os.remove(chart_filename)

        buffer = BytesIO()
        pdf_output = pdf.output(dest='S').encode('latin1', 'replace')  # Replace characters that cannot be encoded
        buffer.write(pdf_output)
        buffer.seek(0)
        return buffer

    generate_button = st.button(f"Generate {document_type}")

    # Initialize the charts_to_generate variable to ensure it's always defined
    if "charts_to_generate" not in st.session_state:
        st.session_state["charts_to_generate"] = []

    # Handle generation and editing
    if generate_button and business_overview:
        with st.spinner(f'Generating your {document_type}...'):
            try:
                # Set token limits based on the selected document length
                if document_length == "Short":
                    max_tokens_per_section = 900  # Total approximately 6000 tokens (6 sections)
                else:  # Long
                    max_tokens_per_section = 2300  # Total approximately 15000 tokens (6 sections)

                # Define sections to generate separately based on document type
                if document_type == "Business Plan":
                    sections = [
                        "Executive Summary",
                        "Company Description",
                        "Market Analysis",
                        "Strategy and Implementation",
                        "Organization and Management Team",
                        "Financial Plan and Projections",
                        "Request for Funding",
                        "Product and Services Description",
                        "SWOT Analysis",
                        "Customer Analysis",
                        "Competitive Analysis",
                        "Marketing Plan",
                        "Operational Plan",
                        "Risk Management",
                        "Exit Strategy",
                        "Appendix"
                    ]
                elif document_type == "Feasibility Study":
                    sections = [
                        "Executive Summary",
                        "Project Description",
                        "Market Feasibility",
                        "Technical Feasibility",
                        "Financial Feasibility",
                        "Economic Feasibility",
                        "Risk Assessment",
                        "Conclusion"
                    ]
                elif document_type == "Application Form":
                    sections = [
                        "Applicant Information",
                        "Business Information",
                        "Funding Request",
                        "Project Details",
                        "Financial Information",
                        "Supporting Documents"
                    ]
                elif document_type == "Pitch Deck":
                    sections = [
                        "Introduction",
                        "Problem Statement",
                        "Solution Overview",
                        "Market Opportunity",
                        "Business Model",
                        "Traction",
                        "Team",
                        "Financial Projections",
                        "Investment Ask",
                        "Closing"
                    ]

                # Context for replacing placeholders
                context = {
                    "Your Name": directors,
                    "Funding Amount": funding_amount,
                    "Grant Amount": grant_amount,
                    "Finance Term": finance_term,
                    "Business Overview": business_overview,
                    "Staffing Compliment": staffing
                }

                document_parts = []
                st.session_state["charts_to_generate"] = []  # Reset charts_to_generate in the session state

                for section in sections:
                    prompt = f"""
                    Generate the {section} of a {document_type.lower()} based on the following inputs:
                    Language: {language}
                    Writing Person: {writing_person}
                    Writing Style: {writing_style}
                    Document Length: {document_length}
                    Template: {template}
                    Business Type: {business_type}
                    BEE Level: {bee_level}
                    Directors/Shareholders: {directors}
                    Staffing Compliment: {staffing}
                    Funding Amount: {funding_amount}
                    Grant Amount: {grant_amount}
                    Finance Term: {finance_term}
                    Business Overview: {business_overview}

                    Ensure that all financial tables, including revenue projections, operating expenses, net profit, and cash flow (if applicable), are calculated accurately based on industry standards, market conditions, and the provided business context. Include all relevant calculations and ensure that all numbers in tables and projections are precise and consistent with the overall document.
                    """

                    response = openai.ChatCompletion.create(
                        model="gpt-4",
                        messages=[
                            {"role": "system", "content": f"You are a helpful assistant with deep knowledge in {document_type.lower()} creation and financial forecasting."},
                            {"role": "user", "content": prompt}
                        ],
                        max_tokens=max_tokens_per_section,  # Adjust based on length selection
                        n=1,
                        stop=None,
                        temperature=0.7
                    )
                    section_text = response['choices'][0]['message']['content']
                    
                    # Replace placeholders with actual user inputs
                    section_text = replace_placeholders(section_text, context)
                    document_parts.append(section_text)
                    
                    # Collect charts to generate at the end
                    if document_type == "Business Plan" and section == "Market Analysis":
                        st.session_state["charts_to_generate"].append(("Market Share Over Time", "Market Analysis: Market Share Over Time"))
                        st.session_state["charts_to_generate"].append(("Competitive Analysis", "Market Analysis: Competitive Analysis"))

                    elif document_type == "Business Plan" and section == "Strategy and Implementation":
                        st.session_state["charts_to_generate"].append(("Milestone Timeline", "Strategy and Implementation: Milestone Timeline"))

                    elif document_type == "Business Plan" and section == "Financial Plan and Projections":
                        st.session_state["charts_to_generate"].append(("Revenue vs Expenses", "Financial Plan: Revenue vs Expenses"))
                        st.session_state["charts_to_generate"].append(("Cash Flow Forecast", "Financial Plan: Cash Flow Forecast"))

                    elif document_type == "Business Plan" and section == "Organization and Management Team":
                        st.session_state["charts_to_generate"].append(("Organizational Structure", "Organization: Organizational Structure"))

                    elif document_type == "Business Plan" and section == "Product and Services Description":
                        st.session_state["charts_to_generate"].append(("Product Development Roadmap", "Product Development: Roadmap"))

                # Combine all parts into a single document
                st.session_state["generated_plan"] = "\n\n".join(document_parts)
                st.session_state["edited_plan"] = st.session_state["generated_plan"]
                
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")

    # Only display editing and downloading options if a plan has been generated
    if st.session_state["generated_plan"]:
        # Allow the user to edit the generated document
        edited_text = st.text_area(
            f"Edit Your {document_type}", 
            value=st.session_state.get("edited_plan", st.session_state["generated_plan"]), 
            height=500, 
            key="edit_area"
        )

        # Ensure the session state updates with every keystroke
        if edited_text != st.session_state["edited_plan"]:
            st.session_state["edited_plan"] = edited_text

        # Display charts at the end of the document
        for chart_type, title in st.session_state["charts_to_generate"]:
            chart_filename = f"{chart_type.replace(' ', '_')}.png"
            save_chart_as_image(chart_type, chart_filename)
            st.image(chart_filename)

        # Options to download the document in various formats
        docx_download = st.download_button(label="Download DOCX", data=convert_to_docx(st.session_state["edited_plan"], st.session_state["charts_to_generate"]), file_name=f"{document_type.replace(' ', '_')}.docx")
        pdf_download = st.download_button(label="Download PDF", data=convert_to_pdf(st.session_state["edited_plan"], st.session_state["charts_to_generate"]), file_name=f"{document_type.replace(' ', '_')}.pdf")

    # Home content here
elif st.session_state.page == "financial_model":
    st.title("Comprehensive Financial Modeling")
    financial_modeling_page()
elif st.session_state.page == "meeting_note_taker":
    st.title("Meeting Note Taker")
    # Display Meeting Note Taker content here
    meeting_note_taker()  # Call the function from your meeting_note_taker.py



